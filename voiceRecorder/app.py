import customtkinter as ctk
import tkinter as tk
from tkinter import messagebox, filedialog
import speech_recognition as sr
import threading
import time
import math
import struct
import pyaudio
import random
import re
# --- NUEVA LIBRERÍA ---
import google.generativeai as genai 
from docx import Document
from fpdf import FPDF

# --- CONFIGURACIÓN DE TEMA JJK ---
ctk.set_appearance_mode("Dark")
ctk.set_default_color_theme("dark-blue")

# --- 🔐 CONFIGURACIÓN GEMINI (Pega tu API KEY aquí) ---
GEMINI_API_KEY = ""  # <--- ¡REEMPLAZA ESTO!
try:
    genai.configure(api_key=GEMINI_API_KEY)
    model = genai.GenerativeModel('gemini-2.5-flash') # O 'gemini-1.5-flash' si quieres más velocidad
except Exception as e:
    print(f"Error configurando Gemini: {e}")

# PALETA DE COLORES
C_BG_MAIN = "#050505"
C_BG_PANEL = "#0F0F0F"
C_ACCENT_GOJO = "#6D28D9"
C_ACCENT_SUKUNA = "#800000"
C_ENERGY_CYAN = "#00B8D4"
C_TEXT_WHITE = "#FAFAFA"
C_TEXT_GRAY = "#616161"
C_BORDER_GRAY = "#333333"

class JJKRecorder(ctk.CTk): 
    def __init__(self):
        super().__init__()
        self.title("Jujutsu Kaisen: Culling Game Recorder")
        self.geometry("750x920")
        self.configure(fg_color=C_BG_MAIN)
        self.resizable(True, True)

        self.is_recording = False
        self.stop_event = threading.Event()
        self.recognizer = sr.Recognizer()
        self.recognizer.dynamic_energy_threshold = True 
        self.full_text = ""
        self.p = pyaudio.PyAudio()
        self.stream = None
        
        self.num_bars = 40
        self.bars = []
        self.noise_gate_threshold = 15 

        self._setup_ui()

    def _setup_ui(self):
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(3, weight=1)

        # 1. ENCABEZADO
        self.header_frame = ctk.CTkFrame(self, fg_color="transparent")
        self.header_frame.grid(row=0, column=0, pady=(35, 15))
        ctk.CTkLabel(self.header_frame, text="EXTENSIÓN DE DOMINIO", font=("Roboto", 28, "bold"), text_color=C_ACCENT_GOJO).pack()
        ctk.CTkLabel(self.header_frame, text="IA GEMINI INTEGRADA", font=("Roboto", 10, "bold"), text_color=C_TEXT_GRAY).pack()

        # 2. SELECTOR
        self.type_doc = ctk.CTkComboBox(
            self, 
            values=["📜 Informe (Carta)", "☠️ Sentencia (Evaluación)", "🩸 Pacto (Formulario)", "🧠 Memoria (Resumen Ejecutivo)"],
            width=280, height=45, font=("Roboto", 13),
            fg_color=C_BG_PANEL, border_color=C_BORDER_GRAY, button_color=C_ACCENT_GOJO,
            dropdown_fg_color=C_BG_PANEL, text_color=C_TEXT_WHITE
        )
        self.type_doc.grid(row=1, column=0, pady=10)
        self.type_doc.set("📜 Informe (Carta)")

        # 3. VISUALIZADOR
        self.vis_frame = ctk.CTkFrame(self, fg_color="transparent", height=100)
        self.vis_frame.grid(row=2, column=0, padx=50, pady=20, sticky="ew")
        
        self.canvas = tk.Canvas(self.vis_frame, bg=C_BG_MAIN, highlightthickness=0, height=100)
        self.canvas.pack(fill="both", expand=True)
        
        bar_w = 12; gap = 4
        total_w = (bar_w + gap) * self.num_bars
        start_x = (650 - total_w) / 2 
        for i in range(self.num_bars):
            x0 = start_x + (i * (bar_w + gap))
            rect = self.canvas.create_rectangle(x0, 95, x0 + bar_w, 95, fill=C_BORDER_GRAY, outline="")
            self.bars.append(rect)

        # 4. TEXTO
        self.text_area = ctk.CTkTextbox(
            self, font=("Consolas", 14), fg_color=C_BG_PANEL, text_color=C_TEXT_WHITE,
            corner_radius=15, border_width=1, border_color=C_BORDER_GRAY, activate_scrollbars=True
        )
        self.text_area.grid(row=3, column=0, padx=50, pady=10, sticky="nsew")

        # 5. CONTROLES
        self.controls_frame = ctk.CTkFrame(self, fg_color="transparent")
        self.controls_frame.grid(row=4, column=0, pady=30)

        self.btn_reset = ctk.CTkButton(self.controls_frame, text="🗑️", width=60, height=60, font=("Arial", 20), fg_color="transparent", hover_color=C_BG_PANEL, border_width=2, border_color=C_BORDER_GRAY, corner_radius=30, command=self.reset_recording)
        self.btn_reset.pack(side="left", padx=25)

        self.btn_record = ctk.CTkButton(self.controls_frame, text="🎙️ INICIAR", width=220, height=60, font=("Roboto", 16, "bold"), fg_color=C_ACCENT_GOJO, hover_color="#5B21B6", corner_radius=30, command=self.toggle_recording)
        self.btn_record.pack(side="left", padx=10)

        self.btn_ai = ctk.CTkButton(self.controls_frame, text="✨ IA", width=60, height=60, font=("Roboto", 14, "bold"), fg_color="transparent", hover_color=C_BG_PANEL, border_width=2, border_color=C_ENERGY_CYAN, text_color=C_ENERGY_CYAN, corner_radius=30, command=self.mahoraga_intelligence)
        self.btn_ai.pack(side="left", padx=25)

        # 6. FOOTER
        self.footer = ctk.CTkFrame(self, fg_color="transparent")
        self.footer.grid(row=5, column=0, pady=(0, 30))
        ctk.CTkButton(self.footer, text="DOCX", width=120, height=35, fg_color="transparent", border_width=1, border_color=C_ENERGY_CYAN, text_color=C_ENERGY_CYAN, hover_color="#00333D", command=lambda: self.save_file("word")).pack(side="left", padx=10)
        ctk.CTkButton(self.footer, text="PDF", width=120, height=35, fg_color="transparent", border_width=1, border_color=C_ACCENT_SUKUNA, text_color="#EF5350", hover_color="#3B0000", command=lambda: self.save_file("pdf")).pack(side="left", padx=10)

        self.lbl_status = ctk.CTkLabel(self, text="Esperando energía...", text_color=C_TEXT_GRAY, font=("Roboto", 10))
        self.lbl_status.grid(row=6, column=0, pady=(0, 10))

    # --- LÓGICA DE VISUALIZACIÓN ---
    def audio_visualizer_loop(self):
        try:
            self.stream = self.p.open(format=pyaudio.paInt16, channels=1, rate=44100, input=True, frames_per_buffer=1024)
            while not self.stop_event.is_set():
                try:
                    data = self.stream.read(1024, exception_on_overflow=False)
                    count = len(data) // 2
                    shorts = struct.unpack("%dh" % count, data)
                    rms = math.sqrt(sum(s**2 for s in shorts) / count)
                    level = 0 if rms < self.noise_gate_threshold else min((rms - self.noise_gate_threshold) / 2, 90)
                    self.update_spectrum(level)
                except: break
        except: pass
        finally:
            if self.stream: self.stream.stop_stream(); self.stream.close()

    def update_spectrum(self, level):
        try:
            base_y = 100
            for bar in self.bars:
                if level == 0:
                    target_h = 2; color = C_BORDER_GRAY
                else:
                    target_h = level * random.uniform(0.8, 1.2)
                    if level < 40: color = C_ENERGY_CYAN
                    elif level < 70: color = C_ACCENT_GOJO
                    else: color = "#FF1744"
                coords = self.canvas.coords(bar)
                self.canvas.coords(bar, coords[0], base_y - target_h, coords[2], base_y)
                self.canvas.itemconfig(bar, fill=color)
        except: pass

    # --- 🧠 LÓGICA MAHORAGA (GEMINI AI REAL) ---
    def mahoraga_intelligence(self):
        raw_text = self.text_area.get("1.0", "end").strip()
        if not raw_text:
            messagebox.showwarning("Error", "No hay texto para adaptar.")
            return

        if GEMINI_API_KEY == "TU_API_KEY_AQUI":
            messagebox.showerror("Falta API Key", "Por favor, abre el código y pon tu API Key de Google Gemini.")
            return

        # Feedback visual
        self.lbl_status.configure(text="Mahoraga: Invocando a Gemini (Procesando)...", text_color=C_ENERGY_CYAN)
        self.btn_ai.configure(state="disabled")
        self.update()

        # Obtener el tipo de formato seleccionado (limpiamos emojis)
        raw_selection = self.type_doc.get()
        # Ejemplo: "📜 Informe (Carta)" -> "Informe (Carta)"
        format_type = raw_selection.split(" ", 1)[1] 

        # Crear el PROMPT para Gemini
        prompt = f"""
        Actúa como un asistente de redacción experto y profesional.
        
        Tengo el siguiente texto crudo dictado por voz (puede tener errores, muletillas o falta de puntuación):
        "{raw_text}"
        
        Tu tarea es reescribirlo completamente dándole el formato de: {format_type}.
        
        Instrucciones específicas:
        1. Corrige ortografía y gramática.
        2. Elimina muletillas (eh, estem, mm).
        3. Usa un tono profesional y formal adecuado para {format_type}.
        4. Si es 'Informe' o 'Carta', incluye estructura de saludo y cuerpo.
        5. Si es 'Evaluación', usa viñetas o puntos clave.
        6. NO agregues explicaciones fuera del texto (como "Aquí tienes tu texto"). Solo dame el documento final.
        """

        # Hilo separado para no congelar la UI mientras Gemini piensa
        threading.Thread(target=self.call_gemini, args=(prompt,), daemon=True).start()

    def call_gemini(self, prompt):
        try:
            response = model.generate_content(prompt)
            final_text = response.text
            
            # Actualizar UI con el texto mágico
            self.text_area.delete("1.0", "end")
            self.text_area.insert("end", final_text)
            self.lbl_status.configure(text="Adaptación completada por Gemini.", text_color=C_TEXT_WHITE)
            self.btn_ai.configure(state="normal")
            
        except Exception as e:
            self.lbl_status.configure(text="Error invocando a Gemini", text_color="#EF5350")
            print(e)
            self.btn_ai.configure(state="normal")
            messagebox.showerror("Error de Conexión", f"Mahoraga no respondió: {e}")

    # --- GRABACIÓN ---
    def toggle_recording(self):
        if not self.is_recording: self.start_recording()
        else: self.stop_recording()

    def start_recording(self):
        self.is_recording = True
        self.stop_event.clear()
        self.btn_record.configure(text="🛑 DETENER", fg_color=C_ACCENT_SUKUNA, hover_color="#B71C1C")
        self.lbl_status.configure(text="🔴 Grabando | Técnica Activa", text_color="#EF5350")
        threading.Thread(target=self.record_loop, daemon=True).start()
        threading.Thread(target=self.audio_visualizer_loop, daemon=True).start()

    def stop_recording(self):
        self.is_recording = False
        self.stop_event.set()
        self.btn_record.configure(text="🎙️ INICIAR", fg_color=C_ACCENT_GOJO, hover_color="#5B21B6")
        self.lbl_status.configure(text="Estado: Pausado", text_color=C_TEXT_GRAY)
        self.update_spectrum(0)

    def record_loop(self):
        with sr.Microphone() as source:
            self.recognizer.adjust_for_ambient_noise(source, duration=0.5)
            while not self.stop_event.is_set():
                try:
                    audio = self.recognizer.listen(source, timeout=5, phrase_time_limit=None)
                    try:
                        text = self.recognizer.recognize_google(audio, language="es-ES")
                        self.full_text += text + " "
                        self.text_area.insert("end", text + " ")
                        self.text_area.see("end")
                    except: pass
                except sr.WaitTimeoutError:
                    if not self.stop_event.is_set():
                        self.after(0, self.ask_silence)
                        break
                except: break

    def ask_silence(self):
        self.stop_recording()
        if messagebox.askyesno("Silencio", "¿Continuar el conjuro?"):
            self.start_recording()

    def reset_recording(self):
        if self.is_recording: self.stop_recording()
        if messagebox.askyesno("Purificar", "¿Borrar todo?"):
            self.text_area.delete("1.0", "end")

    def save_file(self, format):
        content = self.text_area.get("1.0", "end").strip()
        if not content: return
        name = self.type_doc.get().split(" ")[1]

        if format == "word":
            f = filedialog.asksaveasfilename(defaultextension=".docx")
            if f:
                d = Document(); d.add_heading(f"--- {name} ---", 0); d.add_paragraph(content); d.save(f)
                messagebox.showinfo("Hecho", "Word sellado.")
        elif format == "pdf":
            f = filedialog.asksaveasfilename(defaultextension=".pdf")
            if f:
                try:
                    p = FPDF(); p.add_page(); p.set_font("Arial", size=12)
                    p.cell(0, 10, f"--- {name} ---", ln=1, align='C')
                    p.multi_cell(0, 6, content.encode('latin-1', 'replace').decode('latin-1'))
                    p.output(f)
                    messagebox.showinfo("Hecho", "PDF sellado.")
                except: pass

if __name__ == "__main__":
    app = JJKRecorder()
    app.mainloop()